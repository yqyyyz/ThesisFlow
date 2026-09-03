"""下载演示用真实 arXiv 论文（四大主题域 + 低相关折叠样本），并将 2 篇转换为 Word/Markdown。
用法: uv run python scripts/download_demo_papers.py
"""
import re
import sys
import time
import urllib.request
import xml.etree.ElementTree as ET
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent.parent / "data" / "demo_papers"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# (arxiv_id, 主题域, 备注)
PAPERS = [
    ("2307.03172", "长上下文核心", "Lost in the Middle"),
    ("2402.13753", "长上下文核心", "LongRoPE 2048k 上下文扩展"),
    ("2404.06654", "长上下文核心", "RULER 长上下文基准"),
    ("2308.14508", "长上下文核心", "LongBench 多任务长上下文基准"),
    ("2309.17453", "长上下文核心", "Attention Sinks / StreamingLLM"),
    ("2306.15595", "长上下文核心", "位置插值扩展上下文窗口"),
    ("2402.14848", "长上下文核心", "输入长度对推理能力影响"),
    ("2312.10997", "RAG与GraphRAG", "RAG 综述"),
    ("2404.16130", "RAG与GraphRAG", "GraphRAG 局部到全局"),
    ("2310.11511", "RAG与GraphRAG", "Self-RAG 自适应检索"),
    ("2401.15884", "RAG与GraphRAG", "CRAG 检索校正"),
    ("2410.05779", "RAG与GraphRAG", "LightRAG 轻量图结构检索"),
    ("2005.11401", "RAG与GraphRAG", "RAG 原始论文"),
    ("2309.15217", "检索评估与校验", "RAGAS 评估框架"),
    ("2303.08896", "检索评估与校验", "SelfCheckGPT 幻觉检测"),
    ("2305.14251", "检索评估与校验", "FActScore 事实性评分"),
    ("2311.05232", "检索评估与校验", "LLM 幻觉综述"),
    ("2404.13501", "Agent记忆与效率", "Agent 记忆机制综述"),
    ("2404.14294", "Agent记忆与效率", "LLM 高效推理综述（含 KV Cache）"),
    ("2307.11088", "长上下文评测", "L-Eval 长上下文评估"),
    # 低相关样本：主题与长上下文完全无关，用于演示「四维打分 < 2.5 自动折叠」
    ("2205.05897", "低相关样本（折叠演示）", "CAGI 基因组解读评估（基因组学）"),
    ("2302.11439", "低相关样本（折叠演示）", "Gd 析出相对马氏体相变与磁热效应的影响（材料科学）"),
    ("1702.05149", "低相关样本（折叠演示）", "湍流预混火焰曲率效应 DNS 研究（流体力学）→ 转 Markdown"),
    ("2002.07310", "低相关样本（折叠演示）", "近邻矮星系中性氢（天体物理）→ 转 Word"),
]

# 需转换为多格式样本的论文：(arxiv_id, 目标格式)
MULTIFORMAT = {
    "1702.05149": ".md",
    "2002.07310": ".docx",
}

BASE = "http://export.arxiv.org/api/query?id_list={ids}&max_results={n}"


def fetch_meta(ids: list[str]) -> dict[str, dict]:
    import subprocess

    url = BASE.format(ids=",".join(ids), n=len(ids))
    out = subprocess.run(
        ["curl", "-sS", "--max-time", "60", "-A", "ThesisFlow-Demo/0.1", url],
        capture_output=True, text=True, timeout=90,
    )
    ns = {"a": "http://www.w3.org/2005/Atom"}
    root = ET.fromstring(out.stdout)
    result = {}
    for entry in root.findall("a:entry", ns):
        eid = entry.find("a:id", ns).text.strip()
        aid = eid.split("/abs/")[-1].split("v")[0]
        title = " ".join(entry.find("a:title", ns).text.split())
        result[aid] = {"title": title}
    return result


def _existing(arxiv_id: str) -> Path | None:
    hits = sorted(OUT_DIR.glob(f"{arxiv_id}_*"))
    hits = [p for p in hits if p.suffix in {".pdf", ".docx", ".md"}]
    return hits[0] if hits else None


def download(arxiv_id: str, title: str) -> bool:
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in title)[:70].strip("_")
    dest = OUT_DIR / f"{arxiv_id}_{safe}.pdf"
    if _existing(arxiv_id):
        print(f"  [跳过已存在] {arxiv_id}", flush=True)
        return True
    url = f"https://arxiv.org/pdf/{arxiv_id}"
    tmp = dest.with_suffix(".part")
    import subprocess

    try:
        subprocess.run(
            [
                "curl", "-sS", "--max-time", "120", "-L",
                "-A", "Mozilla/5.0",
                "-o", str(tmp), url,
            ],
            check=True,
            timeout=180,
        )
        blob = tmp.read_bytes()
        if len(blob) < 50_000 or not blob.startswith(b"%PDF"):
            print(f"  [失败-内容异常 {len(blob)} bytes] {arxiv_id} {title}", flush=True)
            tmp.unlink(missing_ok=True)
            return False
        tmp.replace(dest)
        print(f"  [OK {len(blob)//1024}KB] {dest.name}", flush=True)
        return True
    except Exception as e:
        tmp.unlink(missing_ok=True)
        print(f"  [失败 {type(e).__name__}] {arxiv_id} {title}: {e}", flush=True)
        return False


SECTION_WORDS = {
    "abstract", "introduction", "background", "related work", "method", "methods",
    "methodology", "experiment", "experiments", "results", "discussion", "conclusion",
    "conclusions", "concluding remarks", "references", "acknowledgements", "appendix",
    "摘要", "引言", "背景", "相关工作", "方法", "实验", "结果", "讨论", "结论", "参考文献", "致谢", "附录",
}


def _looks_like_heading(line: str) -> bool:
    s = line.strip().lower().rstrip(".:")
    if len(line) > 90 or line[-1:] in {".", ",", ";"}:
        return False
    if s in SECTION_WORDS:
        return True
    if re.match(r"^\d+(\.\d+)*\s+\w{2,}", s):
        return True
    return len(line) < 70 and line[0].isdigit() is False and line.isupper() and " " in line


def extract_text_from_pdf(pdf_path: Path) -> tuple[str, list[str]]:
    import fitz  # PyMuPDF

    doc = fitz.open(str(pdf_path))
    title, lines = "", []
    for page in doc:
        text = page.get_text("text")
        if not text.strip():
            continue
        for ln in text.split("\n"):
            ln = ln.strip()
            if not ln:
                continue
            if not title and len(ln) < 200:
                title = ln
            lines.append(ln)
        if len(lines) > 4000:
            break
    return title, lines


def to_markdown(title: str, lines: list[str]) -> str:
    out = [f"# {title}", ""]
    for ln in lines:
        if ln == title:
            continue
        if _looks_like_heading(ln):
            out += ["", f"## {ln}", ""]
        else:
            out.append(ln)
    return "\n".join(out)


def to_docx_file(title: str, lines: list[str], dest: Path) -> None:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.core_properties.title = title
    doc.add_heading(title, level=0)
    for ln in lines:
        if ln == title:
            continue
        if _looks_like_heading(ln):
            doc.add_heading(ln, level=1)
        else:
            doc.add_paragraph(ln)
    doc.save(str(dest))


def convert_multiformat() -> list[str]:
    import fitz  # noqa: F401 仅探测依赖

    converted = []
    for aid, ext in MULTIFORMAT.items():
        pdfs = list(OUT_DIR.glob(f"{aid}_*.pdf"))
        if not pdfs:
            print(f"  [转换跳过] 未找到 {aid} 的 PDF")
            continue
        pdf = pdfs[0]
        title, lines = extract_text_from_pdf(pdf)
        dest = pdf.with_suffix(ext)
        if ext == ".md":
            dest.write_text(to_markdown(title, lines), encoding="utf-8")
        else:
            to_docx_file(title, lines, dest)
        pdf.unlink()
        print(f"  [转换 OK] {pdf.name} -> {dest.name}")
        converted.append(dest.name)
    return converted


def main():
    sys.stdout.reconfigure(line_buffering=True)
    ids = [p[0] for p in PAPERS]
    print("核验 arXiv 元数据…")
    try:
        meta = fetch_meta(ids)
    except Exception as e:
        print(f"元数据接口失败：{e}，直接按列表下载")
        meta = {}
    ok, failed = 0, []
    for aid, domain, note in PAPERS:
        title = meta.get(aid, {}).get("title") or note
        print(f"[{domain}] {aid} {title[:60]}", flush=True)
        if download(aid, title):
            ok += 1
        else:
            failed.append((aid, title))
        time.sleep(1)
    print(f"\n完成：{ok}/{len(PAPERS)} 篇下载成功")
    converted = convert_multiformat()
    if converted:
        print(f"多格式转换：{len(converted)} 篇（{converted}）")
    if failed:
        print("失败清单：")
        for aid, t in failed:
            print(f"  {aid} {t}")
        sys.exit(1)


if __name__ == "__main__":
    main()
